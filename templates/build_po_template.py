#!/usr/bin/env python3
"""Build the Purchase Order Carbone .docx template (MVP9e) — Havenbeheer-branded.

Carbone data root is `d`; renders against `purchase_orders` records.
Run:    /tmp/po_tpl_venv/bin/python templates/build_po_template.py
Output: templates/purchase-order-template.docx

Notes:
- status & currency are select fields; the template-print render expands them to
  {value,label} objects, so use `.label`.
- Carbone date patterns: no comma (read as patternIn) and no spaces (stripped) -> use dashes.
- internal_notes is intentionally NOT rendered (P2).
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
OUT = os.path.join(HERE, "purchase-order-template.docx")
LOGO = os.path.join(HERE, "main-logo.png")

# Brand palette
DARK_GREEN = RGBColor(0x00, 0x4D, 0x38)
LIME = RGBColor(0x50, 0xB8, 0x48)
GREY = RGBColor(0x6B, 0x6B, 0x6B)
LIGHT_BG = "EAF3EC"   # very light green for zebra / boxes
DARK_GREEN_HEX = "004D38"
LIME_HEX = "50B848"

FONT = "Calibri"


# ---------- low-level helpers ----------
def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _set_borders(el_pr, edges, sz, color, val="single"):
    borders = OxmlElement("w:tcBorders") if el_pr.tag.endswith("tcPr") else OxmlElement("w:tblBorders")
    for edge in edges:
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), val)
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    el_pr.append(borders)


def no_borders(table):
    tblPr = table._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        b.append(e)
    tblPr.append(b)


def cell_bottom_border(cell, sz, color):
    tcPr = cell._tc.get_or_add_tcPr()
    _set_borders(tcPr, ("bottom",), sz, color)


def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for name, v in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:w"), str(v))
        el.set(qn("w:type"), "dxa")
        m.append(el)
    tcPr.append(m)


def run(p, text, *, size=10, bold=False, color=None, font=FONT, caps=False):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    if caps:
        r.font.all_caps = True
    return r


def para(container, *, space_after=2, space_before=0, align=None):
    p = container.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.0
    if align is not None:
        p.alignment = align
    return p


# ---------- document ----------
doc = Document()

# tighten default style
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10)

sec = doc.sections[0]
sec.top_margin = Cm(1.6)
sec.bottom_margin = Cm(1.4)
sec.left_margin = Cm(1.8)
sec.right_margin = Cm(1.8)

# === Letterhead: logo (left) + company contact (right) ===
head = doc.add_table(rows=1, cols=2)
no_borders(head)
head.autofit = False
head.columns[0].width = Cm(8)
head.columns[1].width = Cm(9.4)
lcell, rcell = head.rows[0].cells
lcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
rcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

lp = lcell.paragraphs[0]
lp.paragraph_format.space_after = Pt(0)
if os.path.exists(LOGO):
    lp.add_run().add_picture(LOGO, width=Cm(5.6))
else:
    run(lp, "N.V. HAVENBEHEER SURINAME", size=14, bold=True, color=DARK_GREEN)

cp = rcell.paragraphs[0]
cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
cp.paragraph_format.space_after = Pt(0)
run(cp, "N.V. Havenbeheer Suriname\n", size=10, bold=True, color=DARK_GREEN)
for line in ("Havenlaan Zuid 5, Paramaribo, Suriname",
             "+597 404044  ·  planning-services@havenbeheer.sr",
             "havenbeheer.com"):
    rp = para(rcell, space_after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    run(rp, line, size=8.5, color=GREY)

# green rule under the letterhead
rule = doc.add_table(rows=1, cols=1)
no_borders(rule)
rc = rule.rows[0].cells[0]
shade(rc, DARK_GREEN_HEX)
rc.height = Pt(2)
rp = rc.paragraphs[0]
rp.paragraph_format.space_after = Pt(0)
rp.paragraph_format.space_before = Pt(0)
run(rp, "", size=2)

# === Title bar: "PURCHASE ORDER" (left) + PO number (right) ===
para(doc, space_before=10, space_after=0)
title = doc.add_table(rows=1, cols=2)
no_borders(title)
title.columns[0].width = Cm(10)
title.columns[1].width = Cm(7.4)
tl, tr = title.rows[0].cells
tl.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
tr.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
tlp = tl.paragraphs[0]
tlp.paragraph_format.space_after = Pt(0)
run(tlp, "PURCHASE ORDER", size=26, bold=True, color=DARK_GREEN)
trp = tr.paragraphs[0]
trp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
trp.paragraph_format.space_after = Pt(0)
run(trp, "PO Number\n", size=8.5, color=GREY, caps=True)
run(trp, "{d.po_number}", size=15, bold=True, color=DARK_GREEN)

para(doc, space_before=6, space_after=0)

# === Vendor (Supplier) | Deliver to ===
meta = doc.add_table(rows=1, cols=2)
no_borders(meta)
meta.columns[0].width = Cm(8.7)
meta.columns[1].width = Cm(8.7)
vcell, dcell = meta.rows[0].cells
for c in (vcell, dcell):
    set_cell_margins(c, top=60, bottom=80, left=120, right=120)
    shade(c, LIGHT_BG)

def section_heading(cell, label):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run(p, label, size=9, bold=True, color=DARK_GREEN, caps=True)

section_heading(vcell, "Supplier")
run(para(vcell, space_after=0), "{d.supplier.name}", size=11, bold=True)
run(para(vcell, space_after=0), "{d.supplier.address}", size=9.5, color=GREY)
run(para(vcell, space_after=0), "{d.supplier.email}", size=9.5, color=GREY)
run(para(vcell, space_after=0), "{d.supplier.phone}", size=9.5, color=GREY)

section_heading(dcell, "Deliver to")
run(para(dcell, space_after=0), "{d.delivery_address.name}", size=11, bold=True)
run(para(dcell, space_after=0), "{d.delivery_address.address}", size=9.5, color=GREY)

para(doc, space_before=10, space_after=0)

# === Line items ===
cols = ("Description", "Qty", "Unit", "Unit Price", "Line Total")
widths = (Cm(7.4), Cm(1.8), Cm(2.2), Cm(3.0), Cm(3.0))
lt = doc.add_table(rows=1, cols=5)
lt.alignment = WD_TABLE_ALIGNMENT.LEFT
# header
for c, txt, w in zip(lt.rows[0].cells, cols, widths):
    c.width = w
    shade(c, DARK_GREEN_HEX)
    set_cell_margins(c, top=50, bottom=50, left=100, right=100)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if txt != "Description":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, txt, size=9.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), caps=True)

# body row (repeated by Carbone via [i] ... [i+1])
body = lt.add_row().cells
markers = ("{d.lines[i].description}", "{d.lines[i].quantity_ordered}", "{d.lines[i].unit_of_measure.name}",
           "{d.lines[i].unit_price:formatN(2)}", "{d.lines[i].line_total:formatN(2)}")
for c, m, w in zip(body, markers, widths):
    c.width = w
    set_cell_margins(c, top=50, bottom=50, left=100, right=100)
    cell_bottom_border(c, 4, "D7E3DA")
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if m != markers[0]:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, m, size=10)
# loop-end marker row (Carbone drops it)
endrow = lt.add_row().cells
run(endrow[0].paragraphs[0], "{d.lines[i+1].description}", size=1)

para(doc, space_before=8, space_after=0)

# === Totals (right-aligned box) ===
tot = doc.add_table(rows=3, cols=2)
no_borders(tot)
tot.alignment = WD_TABLE_ALIGNMENT.RIGHT
tot.columns[0].width = Cm(5.5)
tot.columns[1].width = Cm(4.0)
rows = (("Order Total (lines)", "{d.currency.label} {d.lines_total:formatN(2)}", False),
        ("Invoice Total", "{d.currency.label} {d.total:formatN(2)}", False),
        ("Invoice Total (USD)", "USD {d.total_usd:formatN(2)}", True))
for ri, (label, marker, emph) in enumerate(rows):
    lc, rc2 = tot.rows[ri].cells
    set_cell_margins(lc, top=40, bottom=40, left=100, right=100)
    set_cell_margins(rc2, top=40, bottom=40, left=100, right=100)
    if emph:
        shade(lc, DARK_GREEN_HEX)
        shade(rc2, DARK_GREEN_HEX)
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    lp.paragraph_format.space_after = Pt(0)
    run(lp, label, size=10, bold=True,
        color=(RGBColor(0xFF, 0xFF, 0xFF) if emph else DARK_GREEN), caps=True)
    rp = rc2.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(0)
    run(rp, marker, size=(12 if emph else 10), bold=True,
        color=(RGBColor(0xFF, 0xFF, 0xFF) if emph else RGBColor(0, 0, 0)))

# === Supplier note (hidden when empty via Carbone conditional) ===
para(doc, space_before=12, space_after=0)
np = para(doc, space_after=0)
run(np, "{d.supplier_note:ifNEM():showBegin}", size=9.5)
run(np, "Note to supplier  ", size=9, bold=True, color=DARK_GREEN, caps=True)
run(np, "{d.supplier_note}", size=10)
run(np, "{d.supplier_note:showEnd}", size=9.5)

# === Footer ===
footer_p = sec.footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(footer_p,
    "N.V. Havenbeheer Suriname  ·  Havenlaan Zuid 5, Paramaribo  ·  havenbeheer.com",
    size=8, color=GREY)

doc.save(OUT)
print("wrote", OUT)
